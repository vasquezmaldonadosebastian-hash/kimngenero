# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, ListFlowable, ListItem, PageBreak
from reportlab.lib.utils import ImageReader
from PIL import Image as PILImage

ROOT = Path(r"C:\Users\Sebastian Vasquez\Documents\GitHub\KimnGenero")
OUT = ROOT / "outputs" / "brief_diseno_docs"
OUT.mkdir(parents=True, exist_ok=True)

imgs = {
    "home": ROOT / "outputs" / "presentacion_kimn_v2" / "capturas" / "01_home_full.png",
    "indicadores": ROOT / "outputs" / "presentacion_kimn_v2" / "capturas" / "05_indicadores_full.png",
    "estado": ROOT / "outputs" / "presentacion_kimn_v2" / "capturas_indicadores_powerbi" / "pagina_completa_real" / "estado_agrupado_pagina_completa_real_v2.png",
    "indicador": ROOT / "outputs" / "presentacion_kimn_v2" / "capturas_indicadores_powerbi" / "pagina_completa_real" / "02VGGE-04_pagina_completa_real_v2.png",
}

brand = {
    "blue": "0176DE",
    "dark": "03122E",
    "ink": "1A0A2E",
    "pale": "E8F2FF",
    "bg": "F5F4F8",
    "gold": "FEC60D",
}

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_page_header(section, text):
    header = section.header
    p = header.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.runs[0]
    run.font.name = 'Montserrat'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(brand['blue'])


def style_doc(doc, title, subtitle=None):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    add_page_header(sec, "KimnGenero")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = 'Montserrat'
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(brand['ink'])
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = 'Inter'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(90,90,90)
    doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Montserrat'
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(brand['ink'])


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Montserrat'
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(brand['blue'])


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(10.5)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = 'Inter'
        r.font.size = Pt(10.5)


def fit_image_width(path, max_width_inches=6.2):
    img = PILImage.open(path)
    width, height = img.size
    ratio = height / width
    return Inches(max_width_inches), Inches(max_width_inches * ratio)


def add_image(doc, path, caption=None, max_width_inches=6.2):
    if not path.exists():
        return
    w, h = fit_image_width(path, max_width_inches)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=w, height=h)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name = 'Inter'
        r.font.size = Pt(8.5)
        r.italic = True
        r.font.color.rgb = RGBColor(100,100,100)


def build_executive_docx():
    doc = Document()
    style_doc(doc, "Brief Ejecutivo de Diseno - KimnGenero", "Version sintetica para compartir con direccion y equipo de diseno")
    add_h1(doc, "Que es KimnGenero")
    add_para(doc, "KimnGenero es un sitio institucional que muestra informacion sobre igualdad de genero por medio de cifras, fichas informativas y visualizaciones. Su objetivo es ordenar informacion relevante, volverla facil de consultar y transmitir confianza publica.")
    add_image(doc, imgs['home'], "Pantalla de inicio")
    add_h1(doc, "Las 3 pantallas mas importantes")
    add_h2(doc, "1. Inicio")
    add_bullets(doc, ["mensaje principal", "cifras destacadas", "acceso a indicadores", "explicacion del observatorio"])
    add_image(doc, imgs['home'], "Inicio")
    add_h2(doc, "2. Listado de Indicadores")
    add_bullets(doc, ["buscador", "filtros", "tarjetas de indicadores", "estados y categorias"])
    add_image(doc, imgs['indicadores'], "Listado de indicadores")
    add_h2(doc, "3. Ficha de Indicador")
    add_bullets(doc, ["titulo del indicador", "explicacion breve", "visualizacion principal", "metodologia y ficha tecnica"])
    add_image(doc, imgs['indicador'], "Ficha de indicador")
    add_h1(doc, "Linea visual actual")
    add_bullets(doc, [
        "Azul principal: #0176DE",
        "Azul oscuro: #03122E",
        "Azul oscuro secundario: #173F8A",
        "Azul claro: #E8F2FF",
        "Fondo claro: #F5F4F8",
        "Amarillo acento: #FEC60D",
        "Tipografias: Montserrat para titulos e Inter para lectura"
    ])
    add_h1(doc, "Prioridades de rediseño")
    add_bullets(doc, [
        "Tarjetas del listado de indicadores",
        "Ficha de indicador",
        "Sistema de color institucional y tematico",
        "Pantallas editoriales"
    ])
    path = OUT / "BRIEF_DISENO_GRAFICO_EJECUTIVO.docx"
    doc.save(path)
    return path


def build_recommendations_docx():
    doc = Document()
    style_doc(doc, "Brief de Diseno - Recomendaciones Visuales", "Lineamientos prioritarios para una propuesta grafica")
    add_h1(doc, "Enfoque general recomendado")
    add_bullets(doc, [
        "institucional pero no rigido",
        "contemporaneo pero no pasajero",
        "claro para lectura de datos",
        "sobrio, confiable y facil de recordar",
    ])
    add_h1(doc, "1. Jerarquia visual")
    add_para(doc, "Hoy el sitio ya tiene buenas bases, pero varias pantallas concentran demasiadas senales visuales en un mismo nivel. El rediseño deberia separar con mas claridad lo principal, lo secundario y lo accesorio.")
    add_image(doc, imgs['home'], "Referencia de jerarquia actual")
    add_h1(doc, "2. Sistema de color")
    add_para(doc, "El azul institucional esta bien instalado. La oportunidad esta en ordenar mejor la relacion entre ese azul y los colores tematicos usados para distinguir grupos de indicadores.")
    add_image(doc, imgs['indicadores'], "Uso actual de colores en el listado")
    add_h1(doc, "3. Tarjetas del listado de indicadores")
    add_para(doc, "Las tarjetas son la pieza mas repetida del sitio y donde hoy conviene invertir mas criterio de diseno. Necesitan una lectura mas limpia, mejor orden de datos y una jerarquia mas refinada.")
    add_image(doc, imgs['indicadores'], "Tarjetas de indicadores")
    add_h1(doc, "4. Ficha de indicador")
    add_para(doc, "La ficha ya tiene una base fuerte, pero la visualizacion principal debe ganar mas protagonismo y la informacion complementaria debe sentirse mejor distribuida.")
    add_image(doc, imgs['indicador'], "Ficha de indicador y visualizacion")
    add_h1(doc, "5. Pantallas editoriales")
    add_para(doc, "Pantallas como Sobre el Modelo, Glosario y Contacto necesitan una lectura mas editorial: mejor ritmo, subtitulos mas claros y una estructura visual mas intencionada.")
    add_h1(doc, "6. Entregables esperados")
    add_bullets(doc, [
        "paleta cromatica refinada",
        "jerarquia tipografica",
        "propuesta de tarjetas de indicadores",
        "propuesta de ficha de indicador",
        "lineamientos para pantallas editoriales",
        "criterio unificado de iconos y apoyos graficos",
    ])
    path = OUT / "BRIEF_DISENO_GRAFICO_RECOMENDACIONES.docx"
    doc.save(path)
    return path


def build_screenshots_docx():
    doc = Document()
    style_doc(doc, "Brief de Diseno - KimnGenero", "Version con screenshots integrados para revision visual")
    add_h1(doc, "Resumen del Proyecto")
    add_para(doc, "KimnGenero es una plataforma institucional dedicada a mostrar informacion sobre igualdad de genero por medio de cifras, visualizaciones y fichas informativas. Su proposito es hacer visible informacion relevante de manera clara, confiable y ordenada.")
    add_image(doc, imgs['home'], "Vista general del sitio")
    add_h1(doc, "Mapa del Sitio")
    add_para(doc, "Actualmente el sitio tiene 8 pantallas principales y 1 pantalla de error de apoyo.")
    add_h2(doc, "1. Inicio")
    add_bullets(doc, ["mensaje principal", "cifras destacadas", "accesos a secciones clave", "presentacion del observatorio"])
    add_image(doc, imgs['home'], "Pantalla de Inicio")
    add_h2(doc, "2. Listado de Indicadores")
    add_bullets(doc, ["buscador", "filtros", "tarjetas de indicadores", "estado y datos clave de cada tarjeta"])
    add_image(doc, imgs['indicadores'], "Pantalla de Listado de Indicadores")
    add_h2(doc, "3. Ficha de Indicador")
    add_bullets(doc, ["encabezado del indicador", "visualizacion principal", "metodologia", "ficha tecnica"])
    add_image(doc, imgs['indicador'], "Pantalla de Ficha de Indicador", max_width_inches=5.8)
    add_h2(doc, "4. Vista General")
    add_para(doc, "Es una pantalla de sintesis con una visualizacion consolidada.")
    add_image(doc, imgs['estado'], "Pantalla de Vista General")
    add_h1(doc, "Linea Visual Actual")
    add_bullets(doc, [
        "Azul principal: #0176DE",
        "Azul oscuro: #03122E",
        "Azul oscuro secundario: #173F8A",
        "Azul claro: #E8F2FF",
        "Fondo gris claro: #F5F4F8",
        "Amarillo acento: #FEC60D",
        "Montserrat para titulos y cifras destacadas",
        "Inter para textos de lectura"
    ])
    add_h1(doc, "Entregables Esperados")
    add_bullets(doc, [
        "propuesta de linea grafica general",
        "paleta refinada",
        "jerarquia tipografica",
        "rediseno de tarjetas de indicadores",
        "rediseno de la ficha de indicador",
        "lineamientos para pantallas editoriales",
        "criterio unificado de iconos y apoyos graficos",
    ])
    path = OUT / "BRIEF_DISENO_GRAFICO_CON_SCREENSHOTS.docx"
    doc.save(path)
    return path


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleBlue', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=22, textColor=colors.HexColor('#1A0A2E'), alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='Subtle', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#666666'), alignment=TA_CENTER, spaceAfter=14))
    styles.add(ParagraphStyle(name='H1Blue', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, textColor=colors.HexColor('#1A0A2E'), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name='H2Blue', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, textColor=colors.HexColor('#0176DE'), spaceBefore=8, spaceAfter=4))
    styles['BodyText'].fontName = 'Helvetica'
    styles['BodyText'].fontSize = 10
    styles['BodyText'].leading = 14
    return styles


def rl_img(path, max_w=16*cm, max_h=18*cm):
    if not path.exists():
        return Spacer(1, 0.1*cm)
    img = PILImage.open(path)
    w, h = img.size
    scale = min(max_w / w, max_h / h)
    return Image(str(path), width=w*scale, height=h*scale)


def build_pdf(path, title, subtitle, blocks):
    styles = pdf_styles()
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.8*cm, leftMargin=1.8*cm, topMargin=1.6*cm, bottomMargin=1.6*cm)
    story = [Paragraph(title, styles['TitleBlue'])]
    if subtitle:
        story.append(Paragraph(subtitle, styles['Subtle']))
    for block in blocks:
        kind = block[0]
        if kind == 'h1':
            story += [Paragraph(block[1], styles['H1Blue'])]
        elif kind == 'h2':
            story += [Paragraph(block[1], styles['H2Blue'])]
        elif kind == 'p':
            story += [Paragraph(block[1], styles['BodyText']), Spacer(1, 0.15*cm)]
        elif kind == 'bullets':
            items = [ListItem(Paragraph(i, styles['BodyText'])) for i in block[1]]
            story += [ListFlowable(items, bulletType='bullet', start='circle', leftIndent=14), Spacer(1, 0.2*cm)]
        elif kind == 'img':
            story += [rl_img(block[1]), Spacer(1, 0.1*cm), Paragraph(block[2], ParagraphStyle('cap', parent=styles['BodyText'], fontSize=8, textColor=colors.HexColor('#666666'), alignment=TA_CENTER)), Spacer(1, 0.35*cm)]
        elif kind == 'pagebreak':
            story += [PageBreak()]
    doc.build(story)


def build_executive_pdf():
    blocks = [
        ('h1','Que es KimnGenero'),
        ('p','KimnGenero es un sitio institucional que muestra informacion sobre igualdad de genero por medio de cifras, fichas informativas y visualizaciones. Su objetivo es ordenar informacion relevante, volverla facil de consultar y transmitir confianza publica.'),
        ('img', imgs['home'], 'Pantalla de inicio'),
        ('h1','Las 3 pantallas mas importantes'),
        ('h2','1. Inicio'), ('bullets',['mensaje principal','cifras destacadas','acceso a indicadores','explicacion del observatorio']),
        ('h2','2. Listado de Indicadores'), ('bullets',['buscador','filtros','tarjetas de indicadores','estados y categorias']), ('img', imgs['indicadores'], 'Listado de indicadores'),
        ('h2','3. Ficha de Indicador'), ('bullets',['titulo del indicador','explicacion breve','visualizacion principal','metodologia y ficha tecnica']), ('img', imgs['indicador'], 'Ficha de indicador'),
        ('h1','Linea visual actual'), ('bullets',['Azul principal: #0176DE','Azul oscuro: #03122E','Azul oscuro secundario: #173F8A','Azul claro: #E8F2FF','Fondo claro: #F5F4F8','Amarillo acento: #FEC60D','Tipografias: Montserrat e Inter'])
    ]
    path = OUT / 'BRIEF_DISENO_GRAFICO_EJECUTIVO.pdf'
    build_pdf(path, 'Brief Ejecutivo de Diseno - KimnGenero', 'Version sintetica para compartir con direccion y equipo de diseno', blocks)
    return path


def build_recommendations_pdf():
    blocks = [
        ('h1','Enfoque general recomendado'), ('bullets',['institucional pero no rigido','contemporaneo pero no pasajero','claro para lectura de datos','sobrio, confiable y facil de recordar']),
        ('h1','1. Jerarquia visual'), ('p','El rediseño deberia separar con mas claridad lo principal, lo secundario y lo accesorio.'), ('img', imgs['home'], 'Referencia de jerarquia actual'),
        ('h1','2. Sistema de color'), ('p','La oportunidad esta en ordenar mejor la relacion entre el azul institucional y los colores tematicos.'), ('img', imgs['indicadores'], 'Uso actual de colores en el listado'),
        ('h1','3. Tarjetas del listado de indicadores'), ('p','Las tarjetas son la pieza mas repetida del sitio y donde hoy conviene invertir mas criterio de diseno.'), ('img', imgs['indicadores'], 'Tarjetas de indicadores'),
        ('h1','4. Ficha de indicador'), ('p','La visualizacion principal debe ganar mas protagonismo y la informacion complementaria debe sentirse mejor distribuida.'), ('img', imgs['indicador'], 'Ficha de indicador y visualizacion'),
        ('h1','Entregables esperados'), ('bullets',['paleta cromatica refinada','jerarquia tipografica','propuesta de tarjetas de indicadores','propuesta de ficha de indicador','lineamientos para pantallas editoriales','criterio unificado de iconos y apoyos graficos'])
    ]
    path = OUT / 'BRIEF_DISENO_GRAFICO_RECOMENDACIONES.pdf'
    build_pdf(path, 'Brief de Diseno - Recomendaciones Visuales', 'Lineamientos prioritarios para una propuesta grafica', blocks)
    return path


def build_screenshots_pdf():
    blocks = [
        ('h1','Resumen del Proyecto'), ('p','KimnGenero es una plataforma institucional dedicada a mostrar informacion sobre igualdad de genero por medio de cifras, visualizaciones y fichas informativas.'), ('img', imgs['home'], 'Vista general del sitio'),
        ('h1','Mapa del Sitio'), ('p','Actualmente el sitio tiene 8 pantallas principales y 1 pantalla de error de apoyo.'),
        ('h2','1. Inicio'), ('bullets',['mensaje principal','cifras destacadas','accesos a secciones clave','presentacion del observatorio']), ('img', imgs['home'], 'Pantalla de Inicio'),
        ('h2','2. Listado de Indicadores'), ('bullets',['buscador','filtros','tarjetas de indicadores','estado y datos clave de cada tarjeta']), ('img', imgs['indicadores'], 'Pantalla de Listado de Indicadores'),
        ('h2','3. Ficha de Indicador'), ('bullets',['encabezado del indicador','visualizacion principal','metodologia','ficha tecnica']), ('img', imgs['indicador'], 'Pantalla de Ficha de Indicador'),
        ('h2','4. Vista General'), ('p','Es una pantalla de sintesis con una visualizacion consolidada.'), ('img', imgs['estado'], 'Pantalla de Vista General'),
        ('h1','Linea Visual Actual'), ('bullets',['Azul principal: #0176DE','Azul oscuro: #03122E','Azul oscuro secundario: #173F8A','Azul claro: #E8F2FF','Fondo gris claro: #F5F4F8','Amarillo acento: #FEC60D','Montserrat para titulos y cifras destacadas','Inter para textos de lectura'])
    ]
    path = OUT / 'BRIEF_DISENO_GRAFICO_CON_SCREENSHOTS.pdf'
    build_pdf(path, 'Brief de Diseno - KimnGenero', 'Version con screenshots integrados para revision visual', blocks)
    return path

paths = [
    build_executive_docx(),
    build_recommendations_docx(),
    build_screenshots_docx(),
    build_executive_pdf(),
    build_recommendations_pdf(),
    build_screenshots_pdf(),
]
for p in paths:
    print(p)


